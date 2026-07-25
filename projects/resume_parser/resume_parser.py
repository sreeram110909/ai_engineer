import os
from pathlib import Path
from time import sleep
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel, field_validator

load_dotenv()

my_api_key = os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("GROQ_API_KEY environment variable is not set.")

client=Groq(api_key=my_api_key)
model = "openai/gpt-oss-120b"


job_description = f"""
Job description

Hi there,


Overview Research Analyst (STEM Domains – AI Evaluation)

This is a remote, short-term AI evaluation research role focused on creating high-quality research and evaluation data for advanced AI systems.

The selected candidates will research scientific and technical information from reliable public sources, such as research papers, government databases, statistical repositories, and technical documentation. They will then create complex, multi-step questions with objectively verifiable answers, validate the information using multiple sources, and document the complete research process.


About the Role

We are engaged in the development of high-quality evaluation data for advanced AI systems, in partnership with a leading global AI research organisation.
We are hiring Research Analysts for our STEM stream to author complex, source-verified research questions across scientific, technological, medical, and financial subject areas.
This is a rigorous, detail-oriented research position. Analysts are responsible for the end-to-end production of research tasks: identifying verifiable facts in public sources, constructing original multi-step research questions, and documenting the verification process to auditable standards.

Key Responsibilities

Conduct in-depth research across publicly available sources, including peer-reviewed publications, regulatory and government databases, statistical repositories, and technical documentation.
Author original, multi-step research questions with a single, objectively verifiable answer, in accordance with internal quality specifications.
Verify all factual claims against multiple independent, publicly accessible sources prior to submission.
Produce complete documentation of the research and verification process for internal review.
Meet defined quality, accuracy, and throughput standards, completing each assignment in full before commencing the next.
Adhere to all project protocols, quality checklists, and confidentiality requirements.

Required Qualifications

Bachelor's degree (completed or final year) in engineering, the sciences, mathematics, medicine, pharmacy, or economics from a recognized institution;
Candidates from premier institutions (IIT, NIT, BITS, IISc, AIIMS, or equivalent) are strongly encouraged to apply.
Demonstrated ability to read and accurately interpret technical material research papers, statistical tables, datasets, and regulatory documents.
Excellent written English, with the ability to draft precise, unambiguous analytical prose.
Strong attention to detail and a disciplined, process-driven working style.
Full-time availability; own laptop and reliable high-speed internet connection.

Preferred Qualifications

Postgraduate degree (M.Tech, M.Sc, MBBS, M.Pharm, MA Economics, or equivalent) in a relevant discipline.
Prior experience in academic research, scientific writing, data analysis, fact-checking, or technical editing.
Participation in olympiads, research fellowships, or comparable merit-based programs has context menu

Offer Details:


Role: Research Analyst STEM Domains (AI Evaluation)

Project Duration: 2 - 3 Months

Location: Remote

Department: Research Operations

Working Hours: 9 AM to 6 PM IST (Minimum 4 Hrs per day)


"""

class jobD(BaseModel):
    role : str
    required_skills : list[str]
    preferred_skills : list[str]
    minimum_experience : float | None
    education_requirements: list[str]
    responsibilities: list[str]
    
jobd_schema = jobD.model_json_schema()

system_prompt = f"""
You are an expert HR assistant.

Your job is to analyze job descriptions and extract
structured information from them.

Return ONLY valid JSON matching this schema:

{jobd_schema}
IMPORTANT:
Do NOT return the schema itself.
Do NOT return fields like "properties", "title" or "type".
Fill the schema with actual information extracted from the job description.

If minimum experience is not mentioned, return null.
If information for a list is missing, return an empty list.
Do not invent information.
"""

user_prompt = f"""
Analyze the following job description:

{job_description}
"""

message_system = {
    "role" : "system",
    "content" : system_prompt
}

message_user = {
    "role" : "user",
    "content" : user_prompt
}

response_format = {
    "type": "json_object",
}

messages = [message_system, message_user]

response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)

answer=response.choices[0].message.content

raw_json = answer

import json
job_data = json.loads(raw_json)

job =  jobD(**job_data)

print(job.minimum_experience)
print(job.education_requirements)


class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = []
    
class Resume(BaseModel):
    name: str
    email: str
    phone: str
    total_experience_years: float | None = None
    skills: list[str] = []
    experiences: list[Experience] = []
    education: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []

    @field_validator("skills", "education", "projects", "certifications", mode="before")
    @classmethod
    def normalize_list_fields(cls, value):
        if value is None:
            return []
        if isinstance(value, str):
            return [value]
        return value
    
resume_schema = Resume.model_json_schema()

from pypdf import PdfReader
from docx import Document

def read_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
            
    return text

def read_docx(file_path):
    document = Document(file_path)
    text = ""
    for para in document.paragraphs:
        text += para.text + "\n"
        
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text

def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    else:
        return None
    
    
resume_folder = Path("resumes")
all_results = []


    
def parse_resume(resume_text):
    system_prompt = f"""
    You are an expert resume parser.

    Extract information from the resume based on its meaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:

    {resume_schema}

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    
    user_prompt = f"""
    Parse the following resume text and extract structured information from it: {resume_text}"""
    
    message_system = {
        "role": "system",
        "content": system_prompt
    }
    message_user = {
        "role" : "user",
        "content" : user_prompt
    }
    messages = [message_system, message_user]
    response_format = {
        "type": "json_object"
    }
    
    response = client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    raw_json = response.choices[0].message.content
    resume_data = json.loads(raw_json)
    resume = Resume(**resume_data)
    return resume

class MatchResult(BaseModel):
    score: float
    details: dict

def final_score(job,resume):
    match_schema = MatchResult.model_json_schema()
    prompt = f"""
    You are an HR recruiter.

    Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    Return JSON matching this schema:

    {match_schema}

    Give me:

    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Whether experience requirement is met
    5. Overall match percentage from 0 to 100
    6. A short final verdict

    Keep the response concise and easy to read.
    """
    message = {
        "role": "user",
        "content": prompt
    }
    messages = [message]
    response_format = {
        "type": "json_object"
    }
    response = client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    data = json.loads(response.choices[0].message.content)
    return MatchResult(**data)

resume_folder = Path("resumes")
all_results=[]
for file_path in resume_folder.iterdir():
    #C:\Users\Pratyush\padho_with_pratyush\week1\day5\resumes\abhay resume new - Abhay Singh.pdf
    if file_path.suffix.lower() not in [".pdf", ".docx"]:
        continue
    print("\nProcessing:", file_path.name)
    resume_text = read_resume(file_path)
    parsed_resume=parse_resume(resume_text) # llm call1
    sleep(5)
    result = final_score(job, parsed_resume) #llm caLL2
    sleep(5)
    print("Score:", result.score)
    all_results.append({
        "name": parsed_resume.name,
        "score": result.score,
        "details": result.details
    })
    
    all_results.sort(
    key=lambda candidate: candidate["score"],
    reverse=True
)
top_2 = all_results[:2]
worst_2 = all_results[-2:]


print("TOP 2 CANDIDATES")
for candidate in top_2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )

    print(candidate["details"])

print("LOWEST 2 CANDIDATES")
for candidate in worst_2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )
    print(candidate["details"])